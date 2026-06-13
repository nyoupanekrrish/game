# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def set_heading(paragraph, text, level=1,
                color=(30, 80, 180), size=18, bold=True):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(30, 80, 180), size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_heading(p, text, level, color, size)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_subheading(text):
    return add_heading(text, level=2, color=(20, 120, 80), size=13)

def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if sub:
        p.paragraph_format.left_indent = Inches(0.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E50B4')
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(10)
            # Alternate row shading
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8EEF8')
                tcPr.append(shd)
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Inches(col_widths[i])
    return table

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("JUMPING ENGINEER")
title_run.bold = True
title_run.font.size = Pt(36)
title_run.font.color.rgb = RGBColor(30, 80, 180)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Complete Technical Game Development Report")
sub_run.font.size = Pt(16)
sub_run.italic = True
sub_run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}\n").font.size = Pt(11)
info_p.add_run("Engine: Python + Pygame  |  File: main.py  |  Lines of Code: ~2,255").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Executive Summary", size=16)
add_divider()
add_body(
    "JUMPING ENGINEER is a 2D endless side-scrolling runner game built entirely in Python "
    "using the Pygame library. The player controls a character who runs automatically from "
    "left to right and must jump over ground obstacles and dodge falling asteroids to "
    "survive as long as possible. The game features increasing difficulty (speed scales with "
    "score), a health/lives system, a cooldown-based double-jump mechanic, a shield ability, "
    "animated sprites, particle effects, a glassmorphism UI, two selectable maps, persistent "
    "best-score saving, and full sound support."
)
add_body("The entire game runs in a single Python file (main.py) with external asset files.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Technology Stack", size=16)
add_divider()

add_subheading("2.1 Core Language & Libraries")
add_table(
    ["Library / Tool", "Version", "Purpose"],
    [
        ["Python", "3.10+", "Primary programming language"],
        ["pygame", "2.x", "Game loop, rendering, input, audio, windowing"],
        ["pygame.gfxdraw", "built-in", "Anti-aliased circle/polygon drawing (hearts, fire particles)"],
        ["math", "built-in", "Trigonometry for animations, arcs, parametric curves"],
        ["random", "built-in", "Randomized obstacle sizes, asteroid positions, particle behavior"],
        ["json", "built-in", "Persistent best score and settings storage"],
        ["os", "built-in", "Asset file existence checks"],
        ["sys", "built-in", "Clean program exit (sys.exit)"],
    ],
    col_widths=[1.8, 1.0, 3.5]
)

doc.add_paragraph()
add_subheading("2.2 Renderer & Display")
add_bullet("Resolution: 1100 × 700 pixels (fixed window)")
add_bullet("Display mode: pygame.DOUBLEBUF for smooth rendering")
add_bullet("Optional FULLSCREEN mode (toggled in settings, persisted to settings.json)")
add_bullet("Target frame rate: 60 FPS (enforced via clock.tick(60))")
add_bullet("Font loading: tries bundled .ttf fonts (Inter, Roboto, Orbitron) then falls back to system fonts (Segoe UI, Arial), then pygame default")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Project Structure", size=16)
add_divider()
add_body("All code lives in a single file. Assets are in the same directory.", bold=True)

add_subheading("3.1 Code Architecture (main.py)")
add_table(
    ["Section / Class", "Lines (approx.)", "Responsibility"],
    [
        ["Constants & Color Palette", "1 – 34", "Screen size, named colors, accent colors for UI"],
        ["Font Loader (load_preferred_font)", "36 – 52", "Cascading font loading with fallback chain"],
        ["UI Helper: draw_outlined_text()", "70 – 90", "Renders text with pixel outline + drop shadow"],
        ["UI Helper: draw_glass_panel()", "93 – 111", "Frosted-glass rounded rectangle panels"],
        ["UI Helper: draw_vignette()", "114 – 122", "Cinematic dark edge overlay"],
        ["class Particle", "124 – 191", "Fire, spark, and smoke particles with physics"],
        ["class FireText", "193 – 263", "Animated fire text effect (used on restart)"],
        ["draw_heart()", "265 – 285", "Parametric heart curve with gfxdraw fill"],
        ["class GlowButton", "287 – 439", "Reusable buttons: image mode + premium gradient text mode"],
        ["class Player", "441 – 713", "Player physics, animation states, double jump, shield, damage"],
        ["class Obstacle", "715 – 807", "Ground obstacles with glow and PNG asset loading"],
        ["class Asteroid", "809 – 877", "Falling asteroid with GIF frame loading and glow halo"],
        ["class Game.__init__", "880 – 967", "Game state, buttons, sounds, background, settings init"],
        ["Game.load_sounds / play_sound", "999 – 1042", "MP3 loading with silent fallbacks"],
        ["Game.generate_background()", "1044 – 1153", "Procedural parallax background (forest or desert)"],
        ["Game.layout_*()", "1155 – 1218", "Deterministic button layout per game state"],
        ["Game.handle_events()", "1220 – 1398", "Full input & state-machine event handling"],
        ["Game.start_game / reset_game", "1400 – 1223", "Game session initialization"],
        ["Game.update()", "1413 – 1576", "Main game logic: physics, spawning, collision, scoring"],
        ["Game.draw_background()", "1578 – 1721", "Sky gradient, parallax scenery, ground gradient"],
        ["Game.draw_ui()", "1722 – 1852", "In-game HUD: score, hearts, speed bar, double-jump badge, FPS"],
        ["Game.draw_start_menu()", "1855 – 1929", "Animated main menu with floating player preview"],
        ["Game.draw_game_over()", "1931 – 1983", "Mission Failed screen with score card"],
        ["Game.draw_tutorial()", "1985 – 2041", "How-to-play screen with controls list"],
        ["Game.draw_pause_overlay()", "2043 – 2064", "Frosted pause card overlay"],
        ["Game.draw_settings_menu()", "2066 – 2094", "Sound, fullscreen, map toggles"],
        ["Game.draw_playing()", "2096 – 2154", "Composes all in-game visuals and buttons"],
        ["Game.draw_countdown()", "2156 – 2211", "3-2-1-GO! countdown with image or text fallback"],
        ["Game.run()", "2243 – 2251", "Main game loop at 60 FPS"],
    ],
    col_widths=[2.2, 1.3, 3.3]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. GAME MECHANICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Game Mechanics", size=16)
add_divider()

add_subheading("4.1 Player")
add_table(
    ["Mechanic", "Details"],
    [
        ["Starting Position", "x=200, y=HEIGHT-GROUND_HEIGHT-120 (ground level)"],
        ["Base Speed", "8 pixels/frame (auto-run, player does not control horizontal movement)"],
        ["Speed Scaling", "+3% every 10 points scored (score // 10 * 0.03 multiplier)"],
        ["Jump", "Velocity -20 px/frame; gravity +1 px/frame² every frame"],
        ["Double Jump", "Available once per airborne session; 10-second cooldown after use"],
        ["Shield", "5-second protection window; 10-second cooldown; blocks obstacle and asteroid hits"],
        ["Lives", "3 hearts maximum; regain 1 heart every 20 points (max 3)"],
        ["Invincibility Frames", "120 frames (2 sec) after taking damage; player flashes every 75ms"],
        ["Death", "3-frame death animation plays; 2-second delay before Game Over screen"],
        ["States", "running → jumping → hurt → dying → dead"],
    ],
    col_widths=[2.0, 4.8]
)

doc.add_paragraph()
add_subheading("4.2 Obstacles (Ground)")
add_bullet("3 obstacle types: obj1 (red), obj2 (brown), obj3 (gray)")
add_bullet("Each obstacle: random width 50–90 px, random height 40–70 px")
add_bullet("Spawn frequency: base 70–120 frames, reduced by score/50, with ×0.7 to ×1.6 random variation")
add_bullet("Move left at player speed; removed when off-screen left")
add_bullet("Pulsing glow effect using sine wave animation")
add_bullet("On collision: either blocked by shield (spark particles) or player loses a life")

doc.add_paragraph()
add_subheading("4.3 Asteroids (Falling)")
add_bullet("Spawn at random X position above screen (y=-80)")
add_bullet("Fall speed: 3–7 px/frame; rotate at random angular speed")
add_bullet("Spawn frequency: every 5–6 seconds (300–360 frames with random variation)")
add_bullet("Glow halo rendered around asteroid body")
add_bullet("Image loaded from obj1.gif; fallback: drawn circle")
add_bullet("Shield destroys asteroid with spark burst; otherwise player loses a life")

doc.add_paragraph()
add_subheading("4.4 Scoring")
add_table(
    ["Event", "Score Change"],
    [
        ["Pass an obstacle (obstacle.x < player.x)", "+1 point"],
        ["Every 10 points", "Play score sound"],
        ["Every 20 points", "+1 heart (max 3)"],
        ["Beat previous best", "Auto-save to best_score.json"],
    ],
    col_widths=[3.0, 3.8]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. VISUAL SYSTEMS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Visual Systems", size=16)
add_divider()

add_subheading("5.1 Particle System")
add_table(
    ["Type", "Size", "Behavior", "Usage"],
    [
        ["fire", "8–15 px", "Rises, gravity 0.15, shrinks, 30–50 frame life", "Hit effects, score milestone, menu ambiance"],
        ["spark", "2–4 px", "Burst spread, gravity 0.2, fast shrink, 20–40 life", "Double jump, shield block, obstacle hit"],
        ["smoke", "10–20 px", "Slow rise, gravity 0.02, grows to 25 px, 60–100 life", "Obstacle collision smoke"],
    ],
    col_widths=[0.8, 0.8, 2.5, 2.7]
)

doc.add_paragraph()
add_subheading("5.2 Background & Parallax")
add_body("Two selectable maps (toggled in Settings, persisted in settings.json):")
add_table(
    ["Map", "Sky Gradient", "Background Elements", "Ground"],
    [
        ["Default (Forest)", "Blue (135,206,235) → lighter blue", "Mountains (green), Trees with foliage, Clouds (12)", "Grass gradient + animated grass blades"],
        ["Desert", "Warm orange-pink gradient", "Sand dunes (warm), Small rocks (14), Pale clouds (6)", "Sand gradient + sand tufts"],
    ],
    col_widths=[1.5, 1.6, 2.2, 1.5]
)
add_body("Parallax speed: mountains 0.05–0.3×, trees 0.3–0.6×, clouds 0.05–0.2× of player speed — creates layered depth illusion.")

doc.add_paragraph()
add_subheading("5.3 UI Visual Style")
add_bullet("Glassmorphism panels: dark semi-transparent fill + white top-strip highlight + colored border")
add_bullet("All text uses draw_outlined_text(): pixel-perfect outline + drop shadow for legibility over any background")
add_bullet("Neon glow buttons: gradient body, animated neon border, scale-on-hover effect")
add_bullet("Vignette: concentric dark rectangles fading inward — applied on menus and game over screen")
add_bullet("Parametric heart shape: 16sin³(t) / 13cos(t)-5cos(2t)-... curve sampled at 8° intervals")
add_bullet("Animated countdown: PNG images scaled with glow halo; text fallback if PNGs missing")

doc.add_paragraph()
add_subheading("5.4 Player Animations")
add_table(
    ["State", "Frames", "Source Files", "Fallback"],
    [
        ["running", "3 frames", "Run1.png, Run2.png, Run3.png", "Blue rectangle"],
        ["jumping", "4 frames", "Jump1.png, Jump2.png, Jump3.png, Jump4.png", "Light-blue rectangle"],
        ["hurt", "1 frame", "hurt.png", "Red rectangle"],
        ["dying", "3 frames", "Die1.png, Die2.png, Die3.png", "Darkening gray rectangles"],
    ],
    col_widths=[1.2, 1.0, 2.5, 2.1]
)
add_body("All sprites scaled to 90×120 px. Shadow ellipse drawn below player, fading with jump height.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. AUDIO SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_heading("6. Audio System", size=16)
add_divider()
add_body("All sounds use pygame.mixer. Missing files are auto-replaced with a 1000-byte silent buffer so the game never crashes.")
add_table(
    ["Sound File", "Event"],
    [
        ["start.mp3", "Game session begins"],
        ["jump.mp3", "Player jumps or double-jumps"],
        ["hurt.mp3", "Player takes damage (non-fatal)"],
        ["die.mp3", "Player death (fatal hit)"],
        ["score.mp3", "Every 10 points scored"],
        ["shield.mp3", "Shield activated"],
        ["click.mp3", "Mouse hover on menu button"],
        ["pick.mp3", "Menu button clicked"],
    ],
    col_widths=[2.0, 4.8]
)
add_body("Master volume: 1.0 when sound is ON, 0.0 when OFF (toggled in settings, set on all Sound objects at once).")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. GAME STATE MACHINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("7. Game State Machine", size=16)
add_divider()
add_table(
    ["State", "Screen Drawn", "Transitions To"],
    [
        ["start_menu", "Animated menu with buttons", "playing, settings, tutorial, exit"],
        ["settings", "Sound/Fullscreen/Map toggles", "start_menu (on BACK)"],
        ["tutorial", "Controls + tips card", "start_menu"],
        ["playing", "Game world + HUD", "game_over (on death), start_menu (pause→menu)"],
        ["game_over", "Mission Failed + score card", "playing (restart), start_menu"],
    ],
    col_widths=[1.3, 2.2, 3.3]
)
add_body("Paused sub-state: game_state stays 'playing' but update() returns early; pause overlay drawn on top.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. INPUT / CONTROLS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("8. Controls", size=16)
add_divider()
add_table(
    ["Input", "Action"],
    [
        ["SPACE / UP / W", "Jump (ground) or Double Jump (in air)"],
        ["S", "Activate Shield"],
        ["P / ESC", "Toggle Pause"],
        ["R (Game Over)", "Restart game"],
        ["M (Game Over)", "Return to Main Menu"],
        ["ENTER (Main Menu)", "Start game"],
        ["Mouse click – JUMP button", "Jump (touch-friendly on-screen button)"],
        ["Mouse click – SHIELD button", "Activate Shield"],
        ["Mouse click – PAUSE button", "Toggle Pause"],
        ["Mouse hover menu buttons", "Glow highlight + hover sound"],
    ],
    col_widths=[2.4, 4.4]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. PERSISTENCE / SAVE DATA
# ══════════════════════════════════════════════════════════════════════════════
add_heading("9. Persistence & Save Data", size=16)
add_divider()
add_table(
    ["File", "Format", "Contents"],
    [
        ["best_score.json", "JSON", '{"best_score": <int>} — updated whenever score exceeds previous best'],
        ["settings.json", "JSON", '{"sound": bool, "fullscreen": bool, "map": 0|1} — saved when leaving Settings'],
        ["game_data.json", "JSON", "Additional game data (present in folder, not actively used in current code)"],
        ["game_stats.json", "JSON", "Statistics file (present in folder, not actively used in current code)"],
    ],
    col_widths=[1.5, 0.8, 4.5]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 10. COMPLETE ASSET LIST
# ══════════════════════════════════════════════════════════════════════════════
add_heading("10. Complete Asset List", size=16)
add_divider()

add_subheading("10.1 Player Sprite Images (PNG)")
add_table(
    ["File", "State", "Size in Code"],
    [
        ["Run1.png", "Running frame 1", "90×120 px"],
        ["Run2.png", "Running frame 2", "90×120 px"],
        ["Run3.png", "Running frame 3", "90×120 px"],
        ["Jump1.png", "Jump frame 1", "90×120 px"],
        ["Jump2.png", "Jump frame 2", "90×120 px"],
        ["Jump3.png", "Jump frame 3", "90×120 px"],
        ["Jump4.png", "Jump frame 4", "90×120 px"],
        ["hurt.png", "Hurt state", "90×120 px"],
        ["Die1.png", "Death frame 1", "90×120 px"],
        ["Die2.png", "Death frame 2", "90×120 px"],
        ["Die3.png", "Death frame 3", "90×120 px"],
    ],
    col_widths=[1.5, 2.5, 1.5]
)

doc.add_paragraph()
add_subheading("10.2 Obstacle & Enemy Images")
add_table(
    ["File", "Type", "Scaled To"],
    [
        ["obj1.png", "Ground obstacle type 1 (red)", "60×70 px"],
        ["obj2.png", "Ground obstacle type 2 (brown)", "70×60 px"],
        ["obj3.png", "Ground obstacle type 3 (gray)", "65×65 px"],
        ["obj1.gif", "Asteroid image (falling enemy)", "60×60 px"],
    ],
    col_widths=[1.5, 2.5, 1.5]
)

doc.add_paragraph()
add_subheading("10.3 UI Button Images (PNG)")
add_table(
    ["File", "Glow Variant", "Used For"],
    [
        ["start.png", "startglow.png", "START button (main menu)"],
        ["exit.png", "exitglow.png", "EXIT button (main menu)"],
        ["settings.png", "settingsglow.png", "SETTINGS button (main menu)"],
        ["howtoplay.png", "howtoplayglow.png", "HOW TO PLAY button (main menu)"],
        ["restart.png", "—", "RESTART button (game over)"],
        ["menu.png", "menuglow.png", "MENU button"],
        ["back.png", "backglow.png", "BACK button"],
        ["sound.png", "—", "SOUND toggle (settings)"],
        ["screen.png", "—", "FULLSCREEN toggle (settings)"],
        ["map.png", "—", "MAP selector (settings)"],
        ["pause.png", "pauseglow.png", "PAUSE button (in-game)"],
        ["resume.png", "resumeglow.png", "RESUME button (pause menu)"],
        ["button.png", "—", "On-screen JUMP button (in-game)"],
        ["button1.png", "—", "On-screen SHIELD button (in-game)"],
        ["title.png", "—", "Game title logo (main menu)"],
    ],
    col_widths=[1.5, 1.5, 3.8]
)

doc.add_paragraph()
add_subheading("10.4 Countdown Images (PNG)")
add_table(
    ["File", "Displayed When"],
    [
        ["3.png", "First second of countdown after game start"],
        ["2.png", "Second second of countdown"],
        ["1.png", "Third second of countdown"],
        ["go.png", "Fourth second — 'GO!' signal"],
    ],
    col_widths=[1.5, 5.3]
)

doc.add_paragraph()
add_subheading("10.5 Audio Files (MP3)")
add_table(
    ["File", "Trigger"],
    [
        ["start.mp3", "New game starts"],
        ["jump.mp3", "Any jump action"],
        ["hurt.mp3", "Non-fatal hit"],
        ["die.mp3", "Fatal hit (death)"],
        ["score.mp3", "Every 10th point"],
        ["shield.mp3", "Shield activated"],
        ["click.mp3", "Button hover"],
        ["pick.mp3", "Button clicked"],
    ],
    col_widths=[2.0, 4.8]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 11. REQUIRED PACKAGES TO RUN
# ══════════════════════════════════════════════════════════════════════════════
add_heading("11. Setup & Requirements", size=16)
add_divider()

add_subheading("11.1 Python Version")
add_bullet("Python 3.10 or higher (3.12 confirmed to work)")
add_bullet("pygame 2.x  — Install: pip install pygame")

add_subheading("11.2 Install Command")
p = doc.add_paragraph()
p.add_run("pip install pygame").font.name = "Courier New"
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(40, 40, 200)

add_subheading("11.3 Run Command")
p = doc.add_paragraph()
p.add_run("python main.py").font.name = "Courier New"
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(40, 40, 200)

add_subheading("11.4 Minimum File Structure")
for f in ["main.py", "Run1.png, Run2.png, Run3.png", "Jump1.png–Jump4.png",
          "Die1.png–Die3.png", "hurt.png", "obj1.png, obj2.png, obj3.png",
          "obj1.gif (asteroid)", "title.png", "start.png + other button PNGs",
          "1.png, 2.png, 3.png, go.png", "*.mp3 audio files (8 files)"]:
    add_bullet(f)
add_body("Note: If any file is missing, the game uses colored rectangle placeholders for sprites and silent buffers for audio — it will still run.", italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 12. KNOWN ISSUES / BUGS IN ORIGINAL CODE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("12. Known Issues & Notes", size=16)
add_divider()
add_table(
    ["Issue", "Location", "Description"],
    [
        ["Duplicate event handler", "handle_events() lines ~1380 & ~1391", "The game_over elif block is written twice — second copy is unreachable."],
        ["Heart overflow", "update() line ~1557", "add_heart() is called every frame the score is exactly a multiple of 20, causing rapid repeated additions until next frame changes score."],
        ["Hurt.jpg vs hurt.png", "load_animations()", "Code tries to load 'hurt.jpg' but folder contains 'hurt.png'. Sprite falls back to red rectangle."],
        ["Score/Speed text overlap", "draw_ui()", "At high speed values, text may overlap card borders on small displays."],
    ],
    col_widths=[1.8, 2.2, 2.8]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
add_divider()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run(
    f"JUMPING ENGINEER — Technical Report  |  Auto-generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
)
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(140, 140, 140)
footer_run.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "JUMPING_ENGINEER_Report.docx"
doc.save(output_path)
print(f"Report saved: {output_path}")
